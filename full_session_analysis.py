"""
Analyse ICT/SMC complète de TOUS les actifs MT5 (Forex, Indices, Métaux, Crypto)
avec synthèse PDF et DOCX.
================================================================================
Récupère l'historique de tous les symboles disponibles, analyse les sessions
Asian/London/New York, et produit une synthèse ICT/SMC pour chacun.
"""

import sys
import os
import time
from datetime import datetime, timezone, timedelta
from typing import List, Dict, Optional, Tuple
from collections import defaultdict
from dataclasses import dataclass, field

# Forcer UTF-8
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from src.logger_config import setup_logging
logger = setup_logging("full_session_analysis", level="INFO")

import MetaTrader5 as mt5

# Importer le connecteur MT5 existant du projet (avec retries, path, login, etc.)
try:
    from src.mt5_connector import MT5Connector
except ImportError:
    MT5Connector = None

# ─── Configuration ──────────────────────────────────────────────────────────
UTC = timezone.utc
PLATFORM_OFFSET = 2  # UTC+2 (EEST/CEST)

ASIAN_START, ASIAN_END = 0, 8
LONDON_START, LONDON_END = 8, 13
NY_START, NY_END = 13, 21

TIMEFRAME_MAP = {
    "M5": mt5.TIMEFRAME_M5, "M15": mt5.TIMEFRAME_M15,
    "H1": mt5.TIMEFRAME_H1, "H4": mt5.TIMEFRAME_H4,
    "D1": mt5.TIMEFRAME_D1, "W1": mt5.TIMEFRAME_W1,
}

FIB_LEVELS_BULL = (1.618, 2.0, 2.618, 3.0, 3.618, 4.0)
FIB_LEVELS_BEAR = (-1.618, -2.0, -2.618, -3.0, -3.618, -4.0)

# Catégorisation des symboles
FOREX_PAIRS = {
    "EURUSD","GBPUSD","USDJPY","USDCHF","AUDUSD","USDCAD","NZDUSD",
    "EURGBP","EURJPY","EURCHF","EURAUD","EURCAD","EURNZD",
    "GBPJPY","GBPCHF","GBPAUD","GBPCAD","GBPNZD",
    "CHFJPY","AUDJPY","CADJPY","NZDJPY",
    "AUDCAD","AUDCHF","AUDNZD",
    "CADCHF","NZDCAD","NZDCHF",
}
METALS = {"XAUUSD","XAGUSD","XAUAUD","XAUCHF","XAUEUR","XAUGBP","XAGCHF","XAGEUR","XAGGBP"}
INDICES = {"US30","NAS100","SPX500","US500","US100","GER30","GER40","UK100","FRA40","JPN225","AUS200","ESP35","EUSTX50","HK50","CHINA50"}
CRYPTO = {"BTCUSD","ETHUSD","BNBUSD","XRPUSD","LTCUSD","SOLUSD","ADAUSD","DOTUSD"}

def categorize(symbol: str) -> str:
    s = symbol.upper().replace(" ", "").replace("-", "").replace("_", "")
    if s in FOREX_PAIRS or (len(s) == 6 and s[:3] in {"EUR","GBP","USD","AUD","NZD","CAD","CHF","JPY"} and s[3:] in {"EUR","GBP","USD","AUD","NZD","CAD","CHF","JPY"}):
        return "Forex"
    for m in {"XAU","XAG","XPT","XPD","GOLD","SILVER"}:
        if m in s: return "Métaux"
    if "BTC" in s or "ETH" in s or "BNB" in s or "XRP" in s or "LTC" in s or "SOL" in s or "ADA" in s or "DOT" in s:
        return "Crypto"
    for idx in {"US30","NAS","SPX","US500","US100","GER","UK","FRA","JPN","AUS","ESP","STOXX","HK","CHINA","DOW","DJI"}:
        if idx in s: return "Indices"
    return "Autre"

# ─── Helpers ────────────────────────────────────────────────────────────────
def bars_to_dicts(rates_raw) -> List[dict]:
    out = []
    for r in rates_raw:
        out.append({
            "time": int(r[0]), "open": float(r[1]),
            "high": float(r[2]), "low": float(r[3]),
            "close": float(r[4]),
            "tick_volume": int(r[5]) if len(r) > 5 else 0,
        })
    return out

def utc_hour(epoch): return (epoch % 86400) // 3600
def utc_day(epoch): return datetime.fromtimestamp(epoch, UTC).strftime("%Y-%m-%d")
def fmt_p(p, digits=5):
    if p is None: return "-"
    if p >= 1000: return f"{p:.2f}"
    if p >= 10: return f"{p:.4f}"
    return f"{p:.5f}"
def pdt_str(epoch):
    if not epoch: return "-"
    platform = datetime.fromtimestamp(epoch, UTC) + timedelta(hours=PLATFORM_OFFSET)
    return platform.strftime("%Y-%m-%d %H:%M")

def bar_sweeps_high(b, level): return b["high"] > level and b["close"] < level
def bar_sweeps_low(b, level): return b["low"] < level and b["close"] > level

# ─── Analyse d'un symbole ──────────────────────────────────────────────────
@dataclass
class SymbolAnalysis:
    symbol: str
    category: str
    current_bid: float = 0.0
    current_ask: float = 0.0
    spread: float = 0.0
    spread_pct: float = 0.0

    # Asian
    asian_high: Optional[float] = None
    asian_low: Optional[float] = None
    asian_range: float = 0.0
    asian_range_pct: float = 0.0
    asian_mid: float = 0.0
    asian_bars_count: int = 0
    ah_swept: bool = False
    al_swept: bool = False
    ah_swept_at: Optional[float] = None
    al_swept_at: Optional[float] = None
    ah_swept_session: str = "-"
    al_swept_session: str = "-"

    # London
    london_high: Optional[float] = None
    london_low: Optional[float] = None
    london_open: Optional[float] = None
    london_close: Optional[float] = None
    london_range: float = 0.0
    london_direction: str = "-"

    # NY
    ny_high: Optional[float] = None
    ny_low: Optional[float] = None
    ny_open: Optional[float] = None
    ny_close: Optional[float] = None
    ny_range: float = 0.0
    ny_direction: str = "-"

    # Daily levels
    pdh: Optional[float] = None
    pdl: Optional[float] = None
    pdh_date: str = "-"
    pdh_swept: bool = False
    pdl_swept: bool = False

    # Weekly levels
    pwh: Optional[float] = None
    pwl: Optional[float] = None
    pwh_week: str = "-"
    pwh_swept: bool = False
    pwl_swept: bool = False

    # Fibonacci
    fib_target: str = "-"
    fib_next_bull: str = "-"
    fib_next_bear: str = "-"
    bull_fib_swept: List[str] = field(default_factory=list)
    bear_fib_swept: List[str] = field(default_factory=list)

    # Direction & Trade
    direction: str = "NEUTRE"
    direction_detail: str = "-"
    trade_action: str = "-"
    trade_entry: Optional[float] = None
    trade_sl: Optional[float] = None
    trade_tp1: Optional[float] = None
    trade_tp2: Optional[float] = None
    trade_rr: Optional[float] = None
    trade_status: str = "-"

    # Narrative
    narratives: List[str] = field(default_factory=list)

    # Raw bars
    h1_bars: List[dict] = field(default_factory=list)
    m15_bars: List[dict] = field(default_factory=list)
    m5_bars: List[dict] = field(default_factory=list)

    # Error
    error: str = ""

def session_for_hour(h_utc: int) -> str:
    if ASIAN_START <= h_utc < ASIAN_END: return "Asian"
    if LONDON_START <= h_utc < LONDON_END: return "London"
    if NY_START <= h_utc < NY_END: return "NY"
    return "Off"

def mt5_call_with_retry(func, *args, max_retries=3, delay=0.3, **kwargs):
    """Appelle une fonction MT5 avec retry en cas d'erreur IPC (-10005)."""
    last_err = None
    for attempt in range(max_retries):
        try:
            result = func(*args, **kwargs)
            if result is not None:
                return result
            err = mt5.last_error()
            if err and err[0] == -10005:  # IPC timeout
                last_err = err
                time.sleep(delay * (attempt + 1))
                continue
            return result  # None mais pas d'erreur IPC
        except Exception as e:
            last_err = str(e)
            time.sleep(delay * (attempt + 1))
    logger.debug("MT5 call failed after %d retries: %s", max_retries, last_err)
    return None

def analyze_symbol(symbol: str, today_str: str) -> SymbolAnalysis:
    a = SymbolAnalysis(symbol=symbol, category=categorize(symbol))

    try:
        info = mt5_call_with_retry(mt5.symbol_info, symbol)
        if info is None:
            a.error = "Non disponible"
            return a

        if not info.visible:
            mt5_call_with_retry(mt5.symbol_select, symbol, True)

        tick = mt5_call_with_retry(mt5.symbol_info_tick, symbol)
        if tick:
            a.current_bid = float(tick.bid)
            a.current_ask = float(tick.ask)
        else:
            # Fallback: utiliser le close de la dernière barre H1 si disponible
            h1_rates = mt5_call_with_retry(mt5.copy_rates_from_pos, symbol, mt5.TIMEFRAME_H1, 0, 2)
            if h1_rates is not None and len(h1_rates) > 0:
                last_h1 = bars_to_dicts(h1_rates)[-1]
                a.current_bid = last_h1["close"]
                a.current_ask = last_h1["close"]
            else:
                a.error = "Pas de tick ni de barre H1"
                return a
        a.spread = a.current_ask - a.current_bid if a.current_ask and a.current_bid else 0
        a.spread_pct = (a.spread / a.current_bid * 100) if a.current_bid and a.current_bid > 0 else 0

        # Récupération des barres
        h1_rates = mt5_call_with_retry(mt5.copy_rates_from_pos, symbol, mt5.TIMEFRAME_H1, 0, 120)
        m15_rates = mt5_call_with_retry(mt5.copy_rates_from_pos, symbol, mt5.TIMEFRAME_M15, 0, 200)
        m5_rates = mt5_call_with_retry(mt5.copy_rates_from_pos, symbol, mt5.TIMEFRAME_M5, 0, 300)

        if h1_rates is not None:
            a.h1_bars = [b for b in bars_to_dicts(h1_rates) if utc_day(b["time"]) == today_str]
        if m15_rates is not None:
            a.m15_bars = [b for b in bars_to_dicts(m15_rates) if utc_day(b["time"]) == today_str]
        if m5_rates is not None:
            a.m5_bars = [b for b in bars_to_dicts(m5_rates) if utc_day(b["time"]) == today_str]

        all_h1 = bars_to_dicts(h1_rates) if h1_rates is not None else []
        all_h1_today = a.h1_bars

        if not all_h1_today:
            a.error = "Pas de barres H1 aujourd'hui"
            return a

        # ── Asian Range ────────────────────────────────────────────────
        asian_bars = [b for b in all_h1_today if ASIAN_START <= utc_hour(b["time"]) < ASIAN_END]

        # Fallback M15
        if len(asian_bars) < 2:
            asian_m15 = [b for b in a.m15_bars if ASIAN_START <= utc_hour(b["time"]) < ASIAN_END]
            if len(asian_m15) >= 2:
                asian_bars = asian_m15  # use M15 for range calc, but note it

        if len(asian_bars) >= 2:
            a.asian_high = max(b["high"] for b in asian_bars)
            a.asian_low = min(b["low"] for b in asian_bars)
            a.asian_range = a.asian_high - a.asian_low
            a.asian_mid = (a.asian_high + a.asian_low) / 2
            a.asian_bars_count = len(asian_bars)
            a.asian_range_pct = (a.asian_range / a.asian_mid * 100) if a.asian_mid > 0 else 0
            last_asian_time = max(b["time"] for b in asian_bars)

            # Détection sweeps post-Asian
            post_bars = [b for b in all_h1 if b["time"] > last_asian_time]
            for b in post_bars:
                if not a.ah_swept and bar_sweeps_high(b, a.asian_high):
                    a.ah_swept = True
                    a.ah_swept_at = b["time"]
                    a.ah_swept_session = session_for_hour(utc_hour(b["time"]))
                if not a.al_swept and bar_sweeps_low(b, a.asian_low):
                    a.al_swept = True
                    a.al_swept_at = b["time"]
                    a.al_swept_session = session_for_hour(utc_hour(b["time"]))
                if a.ah_swept and a.al_swept:
                    break

            # Fibonacci
            range_size = a.asian_range
            if range_size > 0:
                # Next bull targets
                for n in FIB_LEVELS_BULL:
                    lvl = a.asian_high + n * range_size
                    if a.current_bid < lvl:
                        a.fib_next_bull = f"+{n}"
                        break
                if not a.fib_next_bull:
                    a.fib_next_bull = f"+{FIB_LEVELS_BULL[-1]} ✓"

                for n in FIB_LEVELS_BEAR:
                    lvl = a.asian_low + n * range_size
                    if a.current_bid > lvl:
                        a.fib_next_bear = f"{n}"
                        break
                if not a.fib_next_bear:
                    a.fib_next_bear = f"{FIB_LEVELS_BEAR[-1]} ✓"

                # Swept fib levels
                for n in FIB_LEVELS_BULL:
                    lvl = a.asian_high + n * range_size
                    for b in post_bars:
                        if b["high"] > lvl and b["close"] < lvl:
                            a.bull_fib_swept.append(f"+{n}")
                            break
                for n in FIB_LEVELS_BEAR:
                    lvl = a.asian_low + n * range_size
                    for b in post_bars:
                        if b["low"] < lvl and b["close"] > lvl:
                            a.bear_fib_swept.append(f"{n}")
                            break

                # Fib target directionnel
                if a.al_swept and a.current_bid > a.asian_high:
                    a.fib_target = a.fib_next_bull
                elif a.ah_swept and a.current_bid < a.asian_low:
                    a.fib_target = a.fib_next_bear

            # Direction
            midpoint = a.asian_mid
            if a.ah_swept and a.al_swept:
                a.direction = "MIXTE"
                a.direction_detail = "Les deux extrémités sweepées → range invalide"
            elif a.ah_swept:
                if a.current_bid < a.asian_low:
                    a.direction = "BEAR"
                    a.direction_detail = "AH sweepé, prix sous AL → reversal baissier confirmé"
                elif a.current_bid < midpoint:
                    a.direction = "BEARISH"
                    a.direction_detail = "AH sweepé, prix < midpoint → attente cassure AL"
                else:
                    a.direction = "ATTENTE"
                    a.direction_detail = "AH sweepé mais prix > midpoint → attente"
            elif a.al_swept:
                if a.current_bid > a.asian_high:
                    a.direction = "BULL"
                    a.direction_detail = "AL sweepé, prix au-dessus AH → reversal haussier confirmé"
                elif a.current_bid > midpoint:
                    a.direction = "BULLISH"
                    a.direction_detail = "AL sweepé, prix > midpoint → attente cassure AH"
                else:
                    a.direction = "ATTENTE"
                    a.direction_detail = "AL sweepé mais prix < midpoint → attente"
            else:
                if a.current_bid > a.asian_high:
                    a.direction = "BULLISH"
                    a.direction_detail = "Au-dessus AH, pas de sweep → range testé"
                elif a.current_bid < a.asian_low:
                    a.direction = "BEARISH"
                    a.direction_detail = "Sous AL, pas de sweep → range testé"
                else:
                    a.direction = "NEUTRE"
                    a.direction_detail = "Dans le range, pas de sweep → consolidation"

            # Trade idea
            if a.al_swept and a.current_bid > a.asian_high and not a.ah_swept:
                a.trade_action = "BUY"
                a.trade_entry = a.current_bid
                a.trade_sl = a.asian_low - 0.10 * range_size
                a.trade_tp1 = a.asian_high + 1.618 * range_size
                a.trade_tp2 = a.asian_high + 2.618 * range_size
                risk = a.trade_entry - a.trade_sl
                if risk > 0:
                    a.trade_rr = (a.trade_tp1 - a.trade_entry) / risk
                a.trade_status = "Active"
            elif a.ah_swept and a.current_bid < a.asian_low and not a.al_swept:
                a.trade_action = "SELL"
                a.trade_entry = a.current_bid
                a.trade_sl = a.asian_high + 0.10 * range_size
                a.trade_tp1 = a.asian_low - 1.618 * range_size
                a.trade_tp2 = a.asian_low - 2.618 * range_size
                risk = a.trade_sl - a.trade_entry
                if risk > 0:
                    a.trade_rr = (a.trade_entry - a.trade_tp1) / risk
                a.trade_status = "Active"
            else:
                a.trade_action = "-"
                if a.ah_swept and a.al_swept:
                    a.trade_status = "Range invalide"
                elif not a.ah_swept and not a.al_swept:
                    a.trade_status = "Attente sweep"

        # ── London Session ──────────────────────────────────────────────
        london_bars = [b for b in all_h1_today if LONDON_START <= utc_hour(b["time"]) < LONDON_END]
        if london_bars:
            a.london_high = max(b["high"] for b in london_bars)
            a.london_low = min(b["low"] for b in london_bars)
            a.london_open = london_bars[0]["open"]
            a.london_close = london_bars[-1]["close"]
            a.london_range = a.london_high - a.london_low
            a.london_direction = "HAUSSIER" if a.london_close > a.london_open else "BAISSIER"

        # ── NY Session ──────────────────────────────────────────────────
        ny_bars = [b for b in all_h1_today if NY_START <= utc_hour(b["time"]) < NY_END]
        if ny_bars:
            a.ny_high = max(b["high"] for b in ny_bars)
            a.ny_low = min(b["low"] for b in ny_bars)
            a.ny_open = ny_bars[0]["open"]
            a.ny_close = ny_bars[-1]["close"]
            a.ny_range = a.ny_high - a.ny_low
            a.ny_direction = "HAUSSIER" if a.ny_close > a.ny_open else "BAISSIER"

        # ── Daily Levels ────────────────────────────────────────────────
        try:
            d1_rates = mt5_call_with_retry(mt5.copy_rates_from_pos, symbol, mt5.TIMEFRAME_D1, 0, 4)
            if d1_rates is not None and len(d1_rates) >= 2:
                d1_bars = bars_to_dicts(d1_rates)
                d1_bars.sort(key=lambda b: b["time"])
                yesterday = d1_bars[-2]
                today_d1 = d1_bars[-1]
                a.pdh = yesterday["high"]
                a.pdl = yesterday["low"]
                a.pdh_date = utc_day(yesterday["time"])
                today_open = today_d1["time"]
                today_bars = [b for b in all_h1 if b["time"] >= today_open]
                for b in today_bars:
                    if not a.pdh_swept and bar_sweeps_high(b, a.pdh):
                        a.pdh_swept = True
                    if not a.pdl_swept and bar_sweeps_low(b, a.pdl):
                        a.pdl_swept = True
        except Exception:
            pass

        # ── Weekly Levels ───────────────────────────────────────────────
        try:
            w1_rates = mt5_call_with_retry(mt5.copy_rates_from_pos, symbol, mt5.TIMEFRAME_W1, 0, 3)
            if w1_rates is not None and len(w1_rates) >= 2:
                w1_bars = bars_to_dicts(w1_rates)
                w1_bars.sort(key=lambda b: b["time"])
                last_week = w1_bars[-2]
                this_week = w1_bars[-1]
                a.pwh = last_week["high"]
                a.pwl = last_week["low"]
                a.pwh_week = utc_day(last_week["time"])
                this_week_start = this_week["time"]
                this_week_bars = [b for b in all_h1 if b["time"] >= this_week_start]
                for b in this_week_bars:
                    if not a.pwh_swept and bar_sweeps_high(b, a.pwh):
                        a.pwh_swept = True
                    if not a.pwl_swept and bar_sweeps_low(b, a.pwl):
                        a.pwl_swept = True
        except Exception:
            pass

        # ── Narratives ──────────────────────────────────────────────────
        if a.asian_high:
            a.narratives.append(f"Range asiatique: {fmt_p(a.asian_low)} - {fmt_p(a.asian_high)} (range={fmt_p(a.asian_range)}, {a.asian_range_pct:.2f}% du spot)")
            if a.ah_swept and a.al_swept:
                a.narratives.append(f"Les DEUX extrémités sweepées → range invalidation.")
            elif a.ah_swept:
                a.narratives.append(f"AH sweepé en session {a.ah_swept_session} → raid BSL (stops acheteurs).")
            elif a.al_swept:
                a.narratives.append(f"AL sweepé en session {a.al_swept_session} → raid SSL (stops vendeurs).")
            else:
                a.narratives.append("Aucun sweep du range asiatique — range intact.")

        if a.london_high:
            a.narratives.append(f"London: {a.london_direction}, range {fmt_p(a.london_range)}. O={fmt_p(a.london_open)} C={fmt_p(a.london_close)}")
            if a.asian_high and a.london_high > a.asian_high:
                a.narratives.append(f"London a CASSÉ l'Asian High → liquidity grab haussier.")
            if a.asian_low and a.london_low < a.asian_low:
                a.narratives.append(f"London a CASSÉ l'Asian Low → liquidity grab baissier.")

        if a.ny_high:
            a.narratives.append(f"NY: {a.ny_direction}, range {fmt_p(a.ny_range)}. O={fmt_p(a.ny_open)} C={fmt_p(a.ny_close)}")

        if a.pdh:
            pdh_tag = "SWEPT" if a.pdh_swept else "intact"
            pdl_tag = "SWEPT" if a.pdl_swept else "intact"
            a.narratives.append(f"PDH({a.pdh_date}): {fmt_p(a.pdh)} [{pdh_tag}] | PDL: {fmt_p(a.pdl)} [{pdl_tag}]")
            if a.current_bid > a.pdh:
                a.narratives.append("Prix > PDH → contexte daily bullish.")
            elif a.current_bid < a.pdl:
                a.narratives.append("Prix < PDL → contexte daily bearish.")

        if a.pwh:
            pwh_tag = "SWEPT" if a.pwh_swept else "intact"
            pwl_tag = "SWEPT" if a.pwl_swept else "intact"
            a.narratives.append(f"PWH({a.pwh_week}): {fmt_p(a.pwh)} [{pwh_tag}] | PWL: {fmt_p(a.pwl)} [{pwl_tag}]")

        if a.fib_target and a.fib_target != "-":
            a.narratives.append(f"Cible Fibonacci ICT: {a.fib_target}")

    except Exception as e:
        a.error = str(e)
        logger.debug("Erreur analyse %s: %s", symbol, e)

    return a


# ═══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION PDF
# ═══════════════════════════════════════════════════════════════════════════════

def generate_pdf(analyses: List[SymbolAnalysis], today_str: str, output_path: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm, mm
    from reportlab.lib.colors import HexColor, black, white, grey
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable, KeepTogether,
    )
    from reportlab.platypus.flowables import Flowable

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        rightMargin=1.5*cm, leftMargin=1.5*cm,
        topMargin=1.5*cm, bottomMargin=1.5*cm,
    )
    story = []
    styles = getSampleStyleSheet()

    # Custom styles
    title_style = ParagraphStyle("Title2", parent=styles["Title"], fontSize=18, textColor=HexColor("#1a237e"), spaceAfter=6)
    h1_style = ParagraphStyle("H1", parent=styles["Heading1"], fontSize=14, textColor=HexColor("#1a237e"), spaceBefore=12, spaceAfter=4)
    h2_style = ParagraphStyle("H2", parent=styles["Heading2"], fontSize=12, textColor=HexColor("#283593"), spaceBefore=8, spaceAfter=2)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=8, leading=10, spaceAfter=2)
    small_style = ParagraphStyle("Small", parent=styles["Normal"], fontSize=7, leading=9)
    green_style = ParagraphStyle("Green", parent=body_style, textColor=HexColor("#2e7d32"))
    red_style = ParagraphStyle("Red", parent=body_style, textColor=HexColor("#c62828"))
    yellow_style = ParagraphStyle("Yellow", parent=body_style, textColor=HexColor("#f57f17"))
    blue_style = ParagraphStyle("Blue", parent=body_style, textColor=HexColor("#1565c0"))
    header_style = ParagraphStyle("Header", parent=body_style, fontSize=7, fontName="Helvetica-Bold", textColor=white)

    GREEN = HexColor("#2e7d32")
    RED = HexColor("#c62828")
    YELLOW = HexColor("#f57f17")
    BLUE = HexColor("#1565c0")
    GRAY = HexColor("#9e9e9e")
    DARK = HexColor("#1a237e")
    LIGHT_BG = HexColor("#f5f5f5")
    ASIAN_BG = HexColor("#fff8e1")
    LONDON_BG = HexColor("#ffebee")
    NY_BG = HexColor("#e8f5e9")

    now_str = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")

    # ── Page de garde ───────────────────────────────────────────────────
    story.append(Spacer(1, 3*cm))
    story.append(Paragraph("Analyse ICT/SMC Complète", title_style))
    story.append(Paragraph(f"Sessions Asian, London, New York — {today_str}", h1_style))
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph(f"Généré le {now_str}", styles["Normal"]))
    story.append(Spacer(1, 0.5*cm))

    # Stats globales
    n_total = len(analyses)
    n_ok = len([a for a in analyses if not a.error])
    n_trades = len([a for a in analyses if a.trade_action != "-"])
    n_bull = len([a for a in analyses if a.direction == "BULL"])
    n_bear = len([a for a in analyses if a.direction == "BEAR"])
    n_swept = len([a for a in analyses if a.ah_swept or a.al_swept])

    cats = defaultdict(list)
    for a in analyses: cats[a.category].append(a)

    stats_data = [
        ["Symboles analysés", str(n_total)],
        ["Analyses OK", str(n_ok)],
        ["Erreurs", str(n_total - n_ok)],
        ["Setups directionnels", str(n_trades)],
        ["BULL", str(n_bull)],
        ["BEAR", str(n_bear)],
        ["Sweeps détectés", str(n_swept)],
    ]
    for cat in ["Forex", "Indices", "Métaux", "Crypto", "Autre"]:
        if cats[cat]:
            stats_data.append([f"  {cat}", str(len(cats[cat]))])

    stats_table = Table(stats_data, colWidths=[6*cm, 4*cm])
    stats_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), DARK),
        ("TEXTCOLOR", (0, 0), (-1, 0), white),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 1), (-1, -1), LIGHT_BG),
        ("ALIGN", (1, 0), (1, -1), "CENTER"),
        ("GRID", (0, 0), (-1, -1), 0.5, GRAY),
    ]))
    story.append(stats_table)
    story.append(Spacer(1, 0.5*cm))

    # ── Synthèse des meilleurs setups ──────────────────────────────────
    story.append(Paragraph("Meilleurs Setups Directionnels", h1_style))
    story.append(HRFlowable(width="100%", color=GRAY, thickness=0.5))
    story.append(Spacer(1, 0.2*cm))

    traded = [a for a in analyses if a.trade_action != "-" and a.trade_rr is not None]
    traded.sort(key=lambda a: a.trade_rr or 0, reverse=True)

    if traded:
        synth_header = [["Symbole", "Cat.", "Action", "Entry", "SL", "TP1", "RR", "Direction", "Statut"]]
        for a in traded[:50]:
            rr_str = f"{a.trade_rr:.1f}x" if a.trade_rr else "-"
            synth_header.append([
                a.symbol, a.category, a.trade_action,
                fmt_p(a.trade_entry), fmt_p(a.trade_sl), fmt_p(a.trade_tp1),
                rr_str, a.direction, a.trade_status,
            ])
        t = Table(synth_header, colWidths=[2.2*cm, 1.5*cm, 1.2*cm, 2.2*cm, 2.2*cm, 2.2*cm, 1.2*cm, 1.8*cm, 2.5*cm])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), DARK),
            ("TEXTCOLOR", (0, 0), (-1, 0), white),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("BACKGROUND", (0, 1), (-1, -1), LIGHT_BG),
            ("GRID", (0, 0), (-1, -1), 0.3, GRAY),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [LIGHT_BG, white]),
        ]))
        story.append(t)
    else:
        story.append(Paragraph("Aucun setup directionnel actif.", body_style))

    story.append(PageBreak())

    if not analyses:
        story.append(Paragraph("Aucune analyse à afficher.", body_style))
        doc.build(story)
        return output_path

    # ── Détail par symbole ─────────────────────────────────────────────
    for i, a in enumerate(analyses):
        if i > 0 and i % 3 == 0:
            story.append(PageBreak())

        if a.error:
            story.append(Paragraph(f"<b>{a.symbol}</b> ({a.category}) — <font color='red'>ERREUR: {a.error}</font>", body_style))
            continue

        # Symbole header
        dir_colors = {"BULL": "#2e7d32", "BEAR": "#c62828", "BULLISH": "#4caf50",
                      "BEARISH": "#ef5350", "MIXTE": "#f57f17", "ATTENTE": "#ff9800", "NEUTRE": "#9e9e9e"}
        dir_color = dir_colors.get(a.direction, "#9e9e9e")
        trade_color = "#2e7d32" if a.trade_action == "BUY" else ("#c62828" if a.trade_action == "SELL" else "#9e9e9e")

        # Card-like header
        story.append(Paragraph(
            f"<font size='11' color='#1a237e'><b>{a.symbol}</b></font> "
            f"<font size='8' color='#757575'>({a.category})</font>  |  "
            f"<font size='8'>Bid: <b>{fmt_p(a.current_bid)}</b></font>  |  "
            f"<font size='8'>Spread: {a.spread_pct:.4f}%</font>  |  "
            f"<font size='9' color='{dir_color}'><b>{a.direction}</b></font>"
            , h2_style))
        story.append(HRFlowable(width="100%", color=HexColor("#e0e0e0"), thickness=0.3))

        # ── Tableau Asian + Sweeps ──────────────────────────────────────
        ah_sweep_info = f"{'OUI' if a.ah_swept else 'non'}"
        if a.ah_swept:
            ah_sweep_info += f" ({a.ah_swept_session}@{pdt_str(a.ah_swept_at)})"
        al_sweep_info = f"{'OUI' if a.al_swept else 'non'}"
        if a.al_swept:
            al_sweep_info += f" ({a.al_swept_session}@{pdt_str(a.al_swept_at)})"

        asian_data = [
            ["Asian High", fmt_p(a.asian_high), "London High", fmt_p(a.london_high), "NY High", fmt_p(a.ny_high)],
            ["Asian Low", fmt_p(a.asian_low), "London Low", fmt_p(a.london_low), "NY Low", fmt_p(a.ny_low)],
            ["Range", f"{fmt_p(a.asian_range)} ({a.asian_range_pct:.2f}%)",
             "Range", f"{fmt_p(a.london_range)}" if a.london_range else "-",
             "Range", f"{fmt_p(a.ny_range)}" if a.ny_range else "-"],
            ["AH Swept", ah_sweep_info,
             "AL Swept", al_sweep_info,
             "Fib Target", a.fib_target],
        ]
        asian_table = Table(asian_data, colWidths=[3*cm, 3.5*cm, 3*cm, 3.5*cm, 3*cm, 3.5*cm])
        asian_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), ASIAN_BG),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (2, 0), (2, -1), "Helvetica-Bold"),
            ("FONTNAME", (4, 0), (4, -1), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.3, GRAY),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(asian_table)

        # ── PDH/PDL + PWH/PWL ───────────────────────────────────────────
        if a.pdh or a.pwh:
            levels_data = []
            if a.pdh:
                levels_data.append([
                    Paragraph(f"<b>PDH</b> ({a.pdh_date})", small_style),
                    Paragraph(fmt_p(a.pdh), small_style),
                    Paragraph(f"<font color='#c62828'>SWEPT</font>" if a.pdh_swept else "intact", small_style),
                    Paragraph(f"<b>PDL</b> ({a.pdh_date})", small_style),
                    Paragraph(fmt_p(a.pdl), small_style),
                    Paragraph(f"<font color='#2e7d32'>SWEPT</font>" if a.pdl_swept else "intact", small_style),
                ])
            if a.pwh:
                levels_data.append([
                    Paragraph(f"<b>PWH</b> ({a.pwh_week})", small_style),
                    Paragraph(fmt_p(a.pwh), small_style),
                    Paragraph(f"<font color='#c62828'>SWEPT</font>" if a.pwh_swept else "intact", small_style),
                    Paragraph(f"<b>PWL</b> ({a.pwh_week})", small_style),
                    Paragraph(fmt_p(a.pwl), small_style),
                    Paragraph(f"<font color='#2e7d32'>SWEPT</font>" if a.pwl_swept else "intact", small_style),
                ])
            if levels_data:
                levels_table = Table(levels_data, colWidths=[2.5*cm, 2.5*cm, 2*cm, 2.5*cm, 2.5*cm, 2*cm])
                levels_table.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BG),
                    ("FONTSIZE", (0, 0), (-1, -1), 7),
                    ("GRID", (0, 0), (-1, -1), 0.3, GRAY),
                ]))
                story.append(levels_table)

        # ── Trade Idea ──────────────────────────────────────────────────
        if a.trade_action != "-":
            trade_data = [[
                Paragraph(f"<font size='9' color='{trade_color}'><b>{a.trade_action}</b></font>", body_style),
                Paragraph(f"Entry: <b>{fmt_p(a.trade_entry)}</b>", body_style),
                Paragraph(f"SL: <b>{fmt_p(a.trade_sl)}</b>", body_style),
                Paragraph(f"TP1: <b>{fmt_p(a.trade_tp1)}</b>", body_style),
                Paragraph(f"TP2: <b>{fmt_p(a.trade_tp2)}</b>", body_style),
                Paragraph(f"RR: <b>{a.trade_rr:.1f}x</b>" if a.trade_rr else "RR: -", body_style),
            ]]
            trade_table = Table(trade_data)
            trade_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), HexColor("#e8f5e9") if a.trade_action == "BUY" else HexColor("#ffebee")),
                ("GRID", (0, 0), (-1, -1), 0.3, GRAY),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ]))
            story.append(trade_table)

        # ── Narratives ──────────────────────────────────────────────────
        if a.narratives:
            for n in a.narratives[:8]:
                story.append(Paragraph(f"• {n}", body_style))

        story.append(Spacer(1, 0.3*cm))
        story.append(HRFlowable(width="100%", color=HexColor("#eeeeee"), thickness=0.2))

    # ── Build ──────────────────────────────────────────────────────────
    doc.build(story)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# GÉNÉRATION DOCX
# ═══════════════════════════════════════════════════════════════════════════════

def generate_docx(analyses: List[SymbolAnalysis], today_str: str, output_path: str):
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(9)
    style.font.name = "Calibri"

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # ── Page de garde ───────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Analyse ICT/SMC Complète")
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1a, 0x23, 0x7e)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Sessions Asian, London, New York — {today_str}")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x28, 0x35, 0x93)

    now_str = datetime.now(UTC).strftime("%Y-%m-%d %H:%M UTC")
    dt_para = doc.add_paragraph()
    dt_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = dt_para.add_run(f"Généré le {now_str}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x75, 0x75, 0x75)

    doc.add_paragraph()

    # ── Statistiques ────────────────────────────────────────────────────
    doc.add_heading("Résumé Global", level=1)
    n_total = len(analyses)
    n_ok = len([a for a in analyses if not a.error])
    n_trades = len([a for a in analyses if a.trade_action != "-"])
    n_bull = len([a for a in analyses if a.direction == "BULL"])
    n_bear = len([a for a in analyses if a.direction == "BEAR"])
    n_swept = len([a for a in analyses if a.ah_swept or a.al_swept])

    cats = defaultdict(list)
    for a in analyses: cats[a.category].append(a)

    stats = [
        ("Symboles analysés:", str(n_total)),
        ("Analyses OK:", str(n_ok)),
        ("Erreurs:", str(n_total - n_ok)),
        ("Setups directionnels:", str(n_trades)),
        ("  BULL:", str(n_bull)),
        ("  BEAR:", str(n_bear)),
        ("Sweeps détectés:", str(n_swept)),
    ]
    for cat in ["Forex", "Indices", "Métaux", "Crypto", "Autre"]:
        if cats[cat]:
            stats.append((f"  {cat}:", str(len(cats[cat]))))

    stats_table = doc.add_table(rows=len(stats)+1, cols=2)
    stats_table.style = "Light Grid Accent 1"
    stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = stats_table.rows[0].cells
    hdr[0].text = "Métrique"
    hdr[1].text = "Valeur"
    for i, (k, v) in enumerate(stats):
        stats_table.rows[i+1].cells[0].text = k
        stats_table.rows[i+1].cells[1].text = v

    doc.add_paragraph()

    # ── Meilleurs setups ────────────────────────────────────────────────
    doc.add_heading("Meilleurs Setups Directionnels", level=1)
    traded = [a for a in analyses if a.trade_action != "-" and a.trade_rr is not None]
    traded.sort(key=lambda a: a.trade_rr or 0, reverse=True)

    if traded:
        trade_table = doc.add_table(rows=min(len(traded), 50)+1, cols=9)
        trade_table.style = "Light Grid Accent 1"
        trade_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        headers = ["Symbole", "Cat.", "Action", "Entry", "SL", "TP1", "RR", "Direction", "Statut"]
        for j, h in enumerate(headers):
            trade_table.rows[0].cells[j].text = h
            trade_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True

        for i, a in enumerate(traded[:50]):
            vals = [a.symbol, a.category, a.trade_action, fmt_p(a.trade_entry),
                    fmt_p(a.trade_sl), fmt_p(a.trade_tp1),
                    f"{a.trade_rr:.1f}x" if a.trade_rr else "-", a.direction, a.trade_status]
            for j, v in enumerate(vals):
                trade_table.rows[i+1].cells[j].text = v

    doc.add_page_break()

    # ── Détail par symbole ─────────────────────────────────────────────
    doc.add_heading("Analyse Détaillée par Symbole", level=1)

    for a in analyses:
        if a.error:
            doc.add_heading(f"{a.symbol} ({a.category})", level=2)
            doc.add_paragraph(f"ERREUR: {a.error}")
            continue

        # Header
        h = doc.add_heading(f"{a.symbol} ({a.category}) — {a.direction}", level=2)
        doc.add_paragraph(f"Bid: {fmt_p(a.current_bid)} | Spread: {a.spread_pct:.4f}%")

        # Asian
        p = doc.add_paragraph()
        run = p.add_run("Session Asiatique (UTC 00:00-08:00):")
        run.bold = True
        doc.add_paragraph(
            f"Asian High: {fmt_p(a.asian_high)} | Asian Low: {fmt_p(a.asian_low)} | "
            f"Range: {fmt_p(a.asian_range)} ({a.asian_range_pct:.2f}% du spot)"
        )
        sweep_str = []
        if a.ah_swept:
            sweep_str.append(f"AH SWEEP en {a.ah_swept_session}" + (f" @{pdt_str(a.ah_swept_at)}" if a.ah_swept_at else ""))
        if a.al_swept:
            sweep_str.append(f"AL SWEEP en {a.al_swept_session}" + (f" @{pdt_str(a.al_swept_at)}" if a.al_swept_at else ""))
        if sweep_str:
            doc.add_paragraph(" | ".join(sweep_str))
        else:
            doc.add_paragraph("Aucun sweep du range asiatique.")

        # London/NY
        if a.london_high:
            doc.add_paragraph(
                f"London: {a.london_direction}, Open={fmt_p(a.london_open)}, "
                f"High={fmt_p(a.london_high)}, Low={fmt_p(a.london_low)}, "
                f"Close={fmt_p(a.london_close)}, Range={fmt_p(a.london_range)}"
            )
        if a.ny_high:
            doc.add_paragraph(
                f"New York: {a.ny_direction}, Open={fmt_p(a.ny_open)}, "
                f"High={fmt_p(a.ny_high)}, Low={fmt_p(a.ny_low)}, "
                f"Close={fmt_p(a.ny_close)}, Range={fmt_p(a.ny_range)}"
            )

        # Levels
        if a.pdh:
            pdh_s = "SWEPT" if a.pdh_swept else "intact"
            pdl_s = "SWEPT" if a.pdl_swept else "intact"
            doc.add_paragraph(f"PDH ({a.pdh_date}): {fmt_p(a.pdh)} [{pdh_s}] | PDL: {fmt_p(a.pdl)} [{pdl_s}]")
        if a.pwh:
            pwh_s = "SWEPT" if a.pwh_swept else "intact"
            pwl_s = "SWEPT" if a.pwl_swept else "intact"
            doc.add_paragraph(f"PWH ({a.pwh_week}): {fmt_p(a.pwh)} [{pwh_s}] | PWL: {fmt_p(a.pwl)} [{pwl_s}]")

        # Fib
        if a.fib_target != "-":
            doc.add_paragraph(f"Cible Fibonacci ICT: {a.fib_target}")
        if a.bull_fib_swept:
            doc.add_paragraph(f"Fib+ sweepés: {', '.join(a.bull_fib_swept)}")
        if a.bear_fib_swept:
            doc.add_paragraph(f"Fib- sweepés: {', '.join(a.bear_fib_swept)}")

        # Direction
        doc.add_paragraph(f"Biais: {a.direction} — {a.direction_detail}")

        # Trade
        if a.trade_action != "-":
            p = doc.add_paragraph()
            run = p.add_run(f"[TRADE ICT] {a.trade_action}")
            run.bold = True
            doc.add_paragraph(
                f"Entry: {fmt_p(a.trade_entry)} | SL: {fmt_p(a.trade_sl)} | "
                f"TP1: {fmt_p(a.trade_tp1)} | TP2: {fmt_p(a.trade_tp2)} | "
                f"RR: {a.trade_rr:.1f}x" if a.trade_rr else "-"
            )

        # Narratives
        if a.narratives:
            p = doc.add_paragraph()
            run = p.add_run("Analyse narrative:")
            run.bold = True
            for n in a.narratives:
                doc.add_paragraph(f"• {n}")

        doc.add_paragraph()

    doc.save(output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    import argparse
    parser = argparse.ArgumentParser(
        description="Analyse ICT/SMC complète de tous les actifs MT5 avec PDF+DOCX"
    )
    parser.add_argument("--limit", type=int, default=200,
                       help="Nombre max de symboles à analyser (défaut: 200)")
    parser.add_argument("--categories", nargs="*", default=None,
                       help="Catégories à analyser (Forex, Indices, Métaux, Crypto, Autre)")
    args = parser.parse_args()

    now = datetime.now(UTC)
    today_str = now.strftime("%Y-%m-%d")
    platform_now = now + timedelta(hours=PLATFORM_OFFSET)

    print()
    print("=" * 90)
    print("  ANALYSE ICT/SMC COMPLÈTE — TOUS LES ACTIFS MT5")
    print(f"  Date: {today_str} UTC | Plateforme: {platform_now.strftime('%Y-%m-%d %H:%M')} (UTC+{PLATFORM_OFFSET})")
    print("=" * 90)
    print()

    # ── Connexion MT5 ──────────────────────────────────────────────────
    print("[1/5] Connexion à MetaTrader 5...")
    if MT5Connector is not None:
        conn = MT5Connector()
        if not conn.initialize():
            print(f"[ERREUR] MT5 initialize failed: {mt5.last_error()}")
            return 1
    else:
        if not mt5.initialize():
            print(f"[ERREUR] MT5 initialize failed: {mt5.last_error()}")
            return 1

    ti = mt5.terminal_info()
    if ti:
        print(f"  Terminal: {ti.name} (build {ti.build})")
    else:
        print("  [ERREUR] Impossible de récupérer les infos du terminal.")
        mt5.shutdown()
        return 1

    # ── Liste des symboles ──────────────────────────────────────────────
    print()
    print("[2/5] Récupération de tous les symboles disponibles...")
    all_symbols = mt5.symbols_get()
    if not all_symbols:
        print("[ERREUR] Aucun symbole trouvé.")
        mt5.shutdown()
        return 1

    symbols = sorted({s.name for s in all_symbols})
    print(f"  {len(symbols)} symboles disponibles chez le broker.")

    # Filtrer les symboles utiles (pas les synthétiques ou exotiques inutiles)
    # On garde tout mais on priorise les catégories connues
    categories = defaultdict(list)
    for s in symbols:
        categories[categorize(s)].append(s)

    print(f"  Forex:  {len(categories.get('Forex', []))}")
    print(f"  Métaux: {len(categories.get('Métaux', []))}")
    print(f"  Indices:{len(categories.get('Indices', []))}")
    print(f"  Crypto: {len(categories.get('Crypto', []))}")
    print(f"  Autre:  {len(categories.get('Autre', []))}")

    # On garde tout (jusqu'à --limit max pour la performance)
    good_symbols = []
    cat_order = args.categories if args.categories else ["Forex", "Indices", "Métaux", "Crypto", "Autre"]
    for cat in cat_order:
        good_symbols.extend(sorted(categories.get(cat, [])))

    if len(good_symbols) > args.limit:
        print(f"  Limite de {args.limit} symboles atteinte (sur {len(good_symbols)}).")
        good_symbols = good_symbols[:args.limit]

    print(f"  → {len(good_symbols)} symboles à analyser.")
    print()

    # ── Analyse de chaque symbole ──────────────────────────────────────
    print("[3/5] Analyse ICT/SMC de chaque symbole...")
    analyses: List[SymbolAnalysis] = []
    errors = 0

    for i, sym in enumerate(good_symbols):
        if (i + 1) % 10 == 0 or i == 0:
            n_ok_now = len([a for a in analyses if not a.error])
            print(f"  [{i+1}/{len(good_symbols)}] Analyse... ({sym}) — {n_ok_now} OK, {i - n_ok_now} erreurs")
        a = analyze_symbol(sym, today_str)
        analyses.append(a)
        if a.error:
            errors += 1
        time.sleep(0.15)  # Éviter de saturer l'IPC MT5

    n_ok = len(analyses) - errors
    n_trades = len([a for a in analyses if a.trade_action != "-"])
    n_swept = len([a for a in analyses if a.ah_swept or a.al_swept])
    print(f"  → {n_ok} analyses OK, {errors} erreurs.")
    print(f"  → {n_swept} symboles avec sweeps, {n_trades} setups directionnels.")
    print()

    # ── Génération PDF ─────────────────────────────────────────────────
    print("[4/5] Génération du rapport PDF...")
    reports_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "reports")
    os.makedirs(reports_dir, exist_ok=True)
    pdf_path = os.path.join(reports_dir, f"ict_analysis_{today_str}.pdf")
    try:
        generate_pdf(analyses, today_str, pdf_path)
        print(f"  PDF généré: {pdf_path}")
    except Exception as e:
        print(f"  [ERREUR PDF] {e}")
        logger.exception("PDF generation failed")

    # ── Génération DOCX ────────────────────────────────────────────────
    print("[5/5] Génération du rapport DOCX...")
    docx_path = os.path.join(reports_dir, f"ict_analysis_{today_str}.docx")
    try:
        generate_docx(analyses, today_str, docx_path)
        print(f"  DOCX généré: {docx_path}")
    except Exception as e:
        print(f"  [ERREUR DOCX] {e}")
        logger.exception("DOCX generation failed")

    # ── Résumé console ──────────────────────────────────────────────────
    print()
    print("=" * 90)
    print("  SYNTHÈSE")
    print("=" * 90)
    print(f"  Symboles analysés : {n_ok}/{len(analyses)}")
    print(f"  Sweeps détectés   : {n_swept}")
    print(f"  Setups actifs     : {n_trades}")
    print()

    # Top 10 trades
    traded = [a for a in analyses if a.trade_action != "-" and a.trade_rr is not None]
    traded.sort(key=lambda a: a.trade_rr or 0, reverse=True)

    if traded:
        print(f"  {'TOP 10 MEILLEURS SETUPS':-^70}")
        print(f"  {'Symbole':<10} {'Cat.':<8} {'Action':>6} {'Entry':>12} {'SL':>12} {'TP1':>12} {'RR':>6} {'Dir':<10}")
        print(f"  {'-'*80}")
        for a in traded[:10]:
            print(f"  {a.symbol:<10} {a.category:<8} {a.trade_action:>6} {fmt_p(a.trade_entry):>12} "
                  f"{fmt_p(a.trade_sl):>12} {fmt_p(a.trade_tp1):>12} "
                  f"{a.trade_rr:>5.1f}x {a.direction:<10}")
    print()

    # Direction summary
    dirs = defaultdict(list)
    for a in analyses:
        if a.direction in ("BULL", "BEAR"):
            dirs[a.direction].append(a.symbol)

    if dirs.get("BULL"):
        print(f"  BULL  ({len(dirs['BULL'])}): {', '.join(sorted(dirs['BULL'])[:30])}")
    if dirs.get("BEAR"):
        print(f"  BEAR  ({len(dirs['BEAR'])}): {', '.join(sorted(dirs['BEAR'])[:30])}")

    print()
    print(f"  Fichiers générés :")
    print(f"    PDF  → {pdf_path}")
    print(f"    DOCX → {docx_path}")
    print("=" * 90)

    mt5.shutdown()
    return 0


if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        print("\n[STOP] Interrompu.")
        mt5.shutdown()
        sys.exit(130)
    except Exception as e:
        print(f"\n[ERREUR] {e}")
        import traceback
        traceback.print_exc()
        mt5.shutdown()
        sys.exit(1)
